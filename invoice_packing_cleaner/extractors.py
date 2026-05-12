from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import pandas as pd

from invoice_packing_cleaner.pdf_text_tools import pdf_page_to_rows


@dataclass(frozen=True)
class ParsedTable:
    label: str
    file_name: str
    kind: str
    dataframe: pd.DataFrame
    note: str = ""
    source_sheet_name: str = ""


class UnsupportedFileType(ValueError):
    pass


class MissingDependency(RuntimeError):
    pass


def parse_uploaded_file(file_name: str, data: bytes) -> list[ParsedTable]:
    suffix = Path(file_name).suffix.lower()

    if suffix in {".xlsx", ".xls", ".xlsm"}:
        return _parse_excel(file_name, data)
    if suffix == ".csv":
        return _parse_csv(file_name, data)
    if suffix == ".pdf":
        return _parse_pdf(file_name, data)
    if suffix == ".docx":
        return _parse_docx(file_name, data)
    if suffix == ".doc":
        raise UnsupportedFileType("目前不直接讀取 .doc，請先另存成 .docx 或 PDF。")

    raise UnsupportedFileType(f"不支援的檔案格式：{suffix or '無副檔名'}")


def _normalize_cell(value: object) -> str:
    if value is None or pd.isna(value):
        return ""
    text = str(value)
    text = text.replace("\r", " ").replace("\n", " ").replace("\t", " ")
    return " ".join(text.split())


def _normalize_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    normalized = df.fillna("").astype(str)
    return normalized.apply(lambda col: col.map(_normalize_cell))


def _parse_excel(file_name: str, data: bytes) -> list[ParsedTable]:
    buffer = BytesIO(data)
    workbook = pd.ExcelFile(buffer)
    tables: list[ParsedTable] = []

    for sheet_name in workbook.sheet_names:
        buffer.seek(0)
        df = pd.read_excel(
            buffer,
            sheet_name=sheet_name,
            header=None,
            dtype=str,
            keep_default_na=False,
        )
        tables.append(
            ParsedTable(
                label=f"{file_name} / {sheet_name}",
                file_name=file_name,
                kind="excel",
                dataframe=_normalize_dataframe(df),
                source_sheet_name=sheet_name,
            )
        )

    return tables


def _parse_csv(file_name: str, data: bytes) -> list[ParsedTable]:
    last_error: Exception | None = None
    for encoding in ("utf-8-sig", "cp950", "big5", "latin1"):
        try:
            df = pd.read_csv(
                BytesIO(data),
                header=None,
                dtype=str,
                keep_default_na=False,
                encoding=encoding,
            )
            return [
                ParsedTable(
                    label=f"{file_name} / CSV",
                    file_name=file_name,
                    kind="csv",
                    dataframe=_normalize_dataframe(df),
                    note=f"已使用 {encoding} 編碼讀取。",
                    source_sheet_name="CSV",
                )
            ]
        except Exception as exc:  # pragma: no cover - depends on input encoding
            last_error = exc

    raise ValueError(f"CSV 讀取失敗：{last_error}")


def _parse_pdf(file_name: str, data: bytes) -> list[ParsedTable]:
    try:
        import pdfplumber
    except ModuleNotFoundError as exc:  # pragma: no cover - dependency guard
        raise MissingDependency("請先安裝 pdfplumber。") from exc

    tables: list[ParsedTable] = []

    with pdfplumber.open(BytesIO(data)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            extracted_tables = page.extract_tables() or []
            if extracted_tables:
                for table_number, table in enumerate(extracted_tables, start=1):
                    rows = _pad_rows(table)
                    if rows:
                        tables.append(
                            ParsedTable(
                                label=f"{file_name} / PDF page {page_number} table {table_number}",
                                file_name=file_name,
                                kind="pdf-table",
                                dataframe=_normalize_dataframe(pd.DataFrame(rows)),
                                source_sheet_name=f"PDF page {page_number} table {table_number}",
                            )
                        )
                continue

            rows = pdf_page_to_rows(page)
            if rows:
                tables.append(
                    ParsedTable(
                        label=f"{file_name} / PDF page {page_number} text",
                        file_name=file_name,
                        kind="pdf-text",
                        dataframe=_normalize_dataframe(pd.DataFrame(rows)),
                        note="PDF 沒有表格線，已先用文字行讀取；若欄位不準，請改用 OCR 或手動欄位對應。",
                        source_sheet_name=f"PDF page {page_number} text",
                    )
                )

    if not tables:
        tables.extend(_parse_pdf_with_ocr(file_name, data))

    if not tables:
        tables.append(
            ParsedTable(
                label=f"{file_name} / PDF",
                file_name=file_name,
                kind="pdf-empty",
                dataframe=pd.DataFrame(),
                note="沒有讀到可用表格或文字；如果這是掃描 PDF，請確認 OCR 套件 rapidocr、onnxruntime、pypdfium2 已安裝。",
                source_sheet_name="PDF",
            )
        )

    return tables


def _parse_pdf_with_ocr(file_name: str, data: bytes) -> list[ParsedTable]:
    try:
        from invoice_packing_cleaner.ocr_tools import ocr_pdf_to_dataframes
    except Exception as exc:  # pragma: no cover - optional OCR stack
        return [
            ParsedTable(
                label=f"{file_name} / PDF OCR",
                file_name=file_name,
                kind="pdf-ocr-missing",
                dataframe=pd.DataFrame(),
                note=f"掃描 PDF 需要 OCR 套件才能讀取：{exc}",
                source_sheet_name="PDF OCR",
            )
        ]

    try:
        ocr_tables = ocr_pdf_to_dataframes(file_name, data)
    except Exception as exc:  # pragma: no cover - OCR runtime depends on input
        return [
            ParsedTable(
                label=f"{file_name} / PDF OCR",
                file_name=file_name,
                kind="pdf-ocr-error",
                dataframe=pd.DataFrame(),
                note=f"OCR 讀取掃描 PDF 失敗：{exc}",
                source_sheet_name="PDF OCR",
            )
        ]

    return [
        ParsedTable(
            label=f"{file_name} / {sheet_name}",
            file_name=file_name,
            kind="pdf-ocr",
            dataframe=_normalize_dataframe(df),
            note="這頁是掃描 PDF，系統已使用 OCR 轉成可選欄位；請檢查辨識文字是否有少字或誤字。",
            source_sheet_name=sheet_name,
        )
        for sheet_name, df in ocr_tables
        if not df.empty
    ]


def _parse_docx(file_name: str, data: bytes) -> list[ParsedTable]:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:  # pragma: no cover - dependency guard
        raise MissingDependency("請先安裝 python-docx。") from exc

    document = Document(BytesIO(data))
    tables: list[ParsedTable] = []

    for table_number, table in enumerate(document.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        rows = _pad_rows(rows)
        if rows:
            tables.append(
                ParsedTable(
                    label=f"{file_name} / Word table {table_number}",
                    file_name=file_name,
                    kind="word-table",
                    dataframe=_normalize_dataframe(pd.DataFrame(rows)),
                    source_sheet_name=f"Word table {table_number}",
                )
            )

    if tables:
        return tables

    lines = [[paragraph.text] for paragraph in document.paragraphs if paragraph.text.strip()]
    if lines:
        return [
            ParsedTable(
                label=f"{file_name} / Word text",
                file_name=file_name,
                kind="word-text",
                dataframe=_normalize_dataframe(pd.DataFrame(lines)),
                note="Word 沒有表格，已用文字行讀取；請手動確認欄位。",
                source_sheet_name="Word text",
            )
        ]

    return [
        ParsedTable(
            label=f"{file_name} / Word",
            file_name=file_name,
            kind="word-empty",
            dataframe=pd.DataFrame(),
            note="Word 檔沒有讀到表格或文字。",
            source_sheet_name="Word",
        )
    ]


def _pad_rows(rows: list[list[object] | None]) -> list[list[object]]:
    clean_rows = [list(row or []) for row in rows]
    if not clean_rows:
        return []
    max_len = max(len(row) for row in clean_rows)
    return [row + [""] * (max_len - len(row)) for row in clean_rows]
